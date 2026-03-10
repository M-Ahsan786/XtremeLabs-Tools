"""
Flask Web Application - Markdown to Word Converter
Single integrated tool with beautiful UI and Python backend
"""

from flask import Flask, render_template, request, send_file, jsonify
import os
import sys
import tempfile
import logging
from io import BytesIO
import zipfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document
from docxcompose.composer import Composer
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Add parent directory to path to import converter
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# Import converter
try:
    from converter import MarkdownConverter
except ImportError as e:
    print(f"Warning: Could not import converter.py: {e}")
    MarkdownConverter = None

# Setup
app = Flask(__name__, template_folder='templates', static_folder='static')
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max
app.config['TEMPLATES_AUTO_RELOAD'] = True
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Routes
@app.route('/')
def index():
    """Serve the main landing page"""
    return render_template('index.html')

@app.route('/converter')
def converter_page():
    """Serve the Word Converter standalone page"""
    return render_template('converter.html')

@app.route('/feature-merger')
def feature_merger_page():
    """Serve the Feature Page Merger standalone page"""
    return render_template('feature_merger.html')

@app.route('/api/convert', methods=['POST'])
def convert():
    """Convert markdown to Word document"""
    try:
        # Get form data
        markdown_content = request.form.get('markdown', '')
        include_images = request.form.get('includeImages', 'true').lower() == 'true'
        add_watermark = request.form.get('addWatermark', 'true').lower() == 'true'
        add_page_border = request.form.get('addPageBorder', 'true').lower() == 'true'
        add_page_numbers = request.form.get('addPageNumbers', 'true').lower() == 'true'
        add_logo = request.form.get('addLogo', 'false').lower() == 'true'
        filename = request.form.get('filename', 'document.md')

        if not markdown_content.strip():
            return jsonify({'error': 'Markdown content is empty'}), 400

        if MarkdownConverter is None:
            return jsonify({'error': 'Converter module not found'}), 500

        logger.info(f"Converting: {filename} (images: {include_images}, watermark: {add_watermark}, borders: {add_page_border}, page_nums: {add_page_numbers}, logo: {add_logo})")

        # Create temporary markdown file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as temp_md:
            temp_md.write(markdown_content)
            temp_md_path = temp_md.name

        # Create output path
        output_name = filename.replace('.md', '').replace('.markdown', '').replace('.txt', '') + '.docx'
        temp_output_path = tempfile.mktemp(suffix='.docx')

        try:
            # Convert using converter
            converter = MarkdownConverter(
                temp_md_path,
                temp_output_path,
                include_images=include_images,
                add_page_numbers=add_page_numbers,
                add_page_border=add_page_border,
                add_watermark=add_watermark,
                add_logo=add_logo
            )
            success = converter.convert()

            if not success:
                return jsonify({'error': 'Conversion failed'}), 500

            # Read the generated file
            with open(temp_output_path, 'rb') as f:
                file_data = f.read()

            # Clean up
            os.unlink(temp_md_path)
            os.unlink(temp_output_path)

            logger.info(f"✓ Conversion successful: {output_name} ({len(file_data)} bytes)")

            # Return file
            return send_file(
                BytesIO(file_data),
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=output_name
            )

        except Exception as e:
            if os.path.exists(temp_md_path):
                os.unlink(temp_md_path)
            if os.path.exists(temp_output_path):
                os.unlink(temp_output_path)
            logger.error(f"Conversion error: {str(e)}")
            return jsonify({'error': f'Conversion error: {str(e)}'}), 500

    except Exception as e:
        logger.error(f"Request error: {str(e)}")
        return jsonify({'error': f'Request error: {str(e)}'}), 400

@app.route('/api/preview', methods=['POST'])
def preview():
    """Generate HTML preview that matches final conversion output"""
    try:
        markdown_content = request.form.get('markdown', '')
        include_images = request.form.get('includeImages', 'true').lower() == 'true'

        if not markdown_content.strip():
            return jsonify({'error': 'Markdown content is empty'}), 400

        if MarkdownConverter is None:
            return jsonify({'error': 'Converter module not found'}), 500

        # Create a temporary markdown file for the converter
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as temp_md:
            temp_md.write(markdown_content)
            temp_md_path = temp_md.name

        try:
            # Use converter to generate HTML preview
            converter = MarkdownConverter(temp_md_path, '', include_images=include_images)
            html_preview = converter.to_html()

            os.unlink(temp_md_path)

            logger.info("✓ Preview generated successfully")

            return jsonify({
                'html': html_preview,
                'status': 'success'
            })

        except Exception as e:
            import traceback
            if os.path.exists(temp_md_path):
                os.unlink(temp_md_path)
            error_trace = traceback.format_exc()
            logger.error(f"Preview error: {str(e)}\n{error_trace}")
            return jsonify({'error': f'Preview error: {str(e)}'}), 500

    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        logger.error(f"Preview request error: {str(e)}\n{error_trace}")
        return jsonify({'error': f'Preview error: {str(e)}'}), 400

@app.route('/api/health', methods=['GET'])
def health():
    """Health check endpoint"""
    return jsonify({'status': 'ok', 'service': 'Markdown to Word Converter'})

@app.route('/api/merge_feature', methods=['POST'])
def merge_feature():
    """Merge a single 'feature' DOCX into multiple target DOCX files.
    Request should include:
      - featureFile: single .docx file
      - targetFiles: multiple .docx files (one or more)
      - reassignPageNumbers: 'true'/'false' (optional)
    Returns a zip file containing updated documents with '_updated' appended to filenames.
    """
    try:
        if 'featureFile' not in request.files:
            return jsonify({'error': 'featureFile missing'}), 400

        feature_file = request.files['featureFile']
        target_files = request.files.getlist('targetFiles')
        reassign = request.form.get('reassignPageNumbers', 'false').lower() == 'true'

        if not target_files:
            return jsonify({'error': 'No target files uploaded'}), 400

        # Save feature to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as f_feat:
            feature_path = f_feat.name
            feature_file.save(feature_path)

        # Prepare in-memory zip
        mem_zip = BytesIO()
        z = zipfile.ZipFile(mem_zip, mode='w', compression=zipfile.ZIP_DEFLATED)

        temp_paths_to_clean = [feature_path]

        for uploaded in target_files:
            original_name = uploaded.filename or 'document.docx'
            # save uploaded target
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as f_t:
                target_path = f_t.name
                uploaded.save(target_path)
            temp_paths_to_clean.append(target_path)
            try:
                from docx.enum.section import WD_SECTION

                # Load target as the base matrix (so its formatting stays 100% native)
                target_as_base = Document(target_path)
                
                # Load feature and add a section break to isolate it
                feature = Document(feature_path)
                feature.add_section(WD_SECTION.NEW_PAGE)
                
                # Insert the feature page at top of the target document
                composer = Composer(target_as_base)
                # Setting remove_property_fields=False preserves drawing shapes 
                composer.insert(0, feature, remove_property_fields=False)
                merged_doc = composer.doc
                
                # Forcefully strip page borders from the feature page (Section 0)
                # so it doesn't inherit the main target document's borders
                if len(merged_doc.sections) > 0:
                    sectPr = merged_doc.sections[0]._sectPr
                    pgBorders = sectPr.find(qn('w:pgBorders'))
                    if pgBorders is not None:
                        sectPr.remove(pgBorders)

                # Optionally reassign page numbers (centered footer)
                if reassign:
                    for section in merged_doc.sections:
                        footer = section.footer
                        # ensure a footer paragraph exists
                        if footer.paragraphs:
                            footer_para = footer.paragraphs[0]
                        else:
                            footer_para = footer.add_paragraph()
                        footer_para.text = ''
                        footer_para.alignment = 1  # CENTER

                        run = footer_para.add_run()
                        fldChar1 = OxmlElement('w:fldChar')
                        fldChar1.set(qn('w:fldCharType'), 'begin')
                        instrText = OxmlElement('w:instrText')
                        instrText.set(qn('xml:space'), 'preserve')
                        instrText.text = 'PAGE'
                        fldChar2 = OxmlElement('w:fldChar')
                        fldChar2.set(qn('w:fldCharType'), 'end')
                        run._r.append(fldChar1)
                        run._r.append(instrText)
                        run._r.append(fldChar2)

                # Save merged doc to temp file
                merged_path = tempfile.mktemp(suffix='.docx')
                merged_doc.save(merged_path)
                temp_paths_to_clean.append(merged_path)

                # Prepare archive name (append _updated before extension)
                name_root, _ = os.path.splitext(original_name)
                arcname = f"{name_root}_updated.docx"
                z.write(merged_path, arcname=arcname)

            except Exception as e:
                import traceback
                error_msg = f"Failed to merge {original_name}:\n{str(e)}\n{traceback.format_exc()}"
                logger.error(error_msg)
                
                # Write error to a text file in the zip so user sees why it failed
                name_root, _ = os.path.splitext(original_name)
                z.writestr(f"ERROR_{name_root}.txt", error_msg)
                continue

        # If zip is completely empty (no merged files and no error files), add a generic error
        if not z.filelist:
            z.writestr("ERROR.txt", "No files could be merged. Please check your uploaded documents.")

        z.close()
        mem_zip.seek(0)

        # cleanup temp files
        for p in temp_paths_to_clean:
            try:
                if os.path.exists(p):
                    os.unlink(p)
            except Exception:
                pass

        return send_file(
            BytesIO(mem_zip.getvalue()),
            mimetype='application/zip',
            as_attachment=True,
            download_name='merged_files.zip'
        )

    except Exception as e:
        logger.error(f"Merge feature error: {e}")
        return jsonify({'error': f'Merge failed: {str(e)}'}), 500

def _convert_single_file(markdown_content, filename, include_images):
    """Helper function to convert a single file"""
    output_name = filename.replace('.md', '').replace('.markdown', '').replace('.txt', '') + '.docx'
    temp_output_path = tempfile.mktemp(suffix='.docx')

    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as temp_md:
        temp_md.write(markdown_content)
        temp_md_path = temp_md.name

    try:
        converter = MarkdownConverter(temp_md_path, temp_output_path, include_images=include_images)
        success = converter.convert()

        if not success:
            raise Exception("Conversion failed")

        with open(temp_output_path, 'rb') as f:
            file_data = f.read()

        return (output_name, file_data, None)
    except Exception as e:
        logger.error(f"Failed to convert {filename}: {str(e)}")
        return (output_name, None, str(e))
    finally:
        if os.path.exists(temp_md_path):
            os.unlink(temp_md_path)
        if os.path.exists(temp_output_path):
            os.unlink(temp_output_path)

@app.route('/api/batch-convert', methods=['POST'])
def batch_convert():
    """Convert multiple markdown files concurrently"""
    try:
        files_data = request.form.getlist('files')
        filenames = request.form.getlist('filenames')
        include_images = request.form.get('includeImages', 'true').lower() == 'true'

        if not files_data or len(files_data) == 0:
            return jsonify({'error': 'No files provided'}), 400

        if MarkdownConverter is None:
            return jsonify({'error': 'Converter module not found'}), 500

        logger.info(f"Batch converting {len(files_data)} files (include_images: {include_images})")

        # Process files concurrently
        results = []
        with ThreadPoolExecutor(max_workers=3) as executor:
            futures = []
            for content, filename in zip(files_data, filenames):
                future = executor.submit(_convert_single_file, content, filename, include_images)
                futures.append(future)

            for future in as_completed(futures):
                try:
                    output_name, file_data, error = future.result()
                    if error:
                        results.append({'name': output_name, 'error': error})
                    else:
                        results.append({'name': output_name, 'data': file_data})
                except Exception as e:
                    logger.error(f"Batch conversion error: {str(e)}")
                    results.append({'error': str(e)})

        # Create ZIP file with all converted documents
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            for result in results:
                if 'data' in result and result['data']:
                    zf.writestr(result['name'], result['data'])

        zip_buffer.seek(0)

        logger.info(f"✓ Batch conversion complete: {len(results)} files")

        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name='converted_documents.zip'
        )

    except Exception as e:
        logger.error(f"Batch request error: {str(e)}")
        return jsonify({'error': f'Batch conversion error: {str(e)}'}), 500

# Error handlers
@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': 'Not found'}), 404

@app.errorhandler(500)
def server_error(error):
    return jsonify({'error': 'Server error'}), 500

if __name__ == '__main__':
    logger.info("=" * 60)
    logger.info("Starting Markdown to Word Converter Web Application")
    logger.info("=" * 60)
    logger.info("📂 Flask: Initialized successfully")
    logger.info("🎨 UI: Beautiful glass-morphism interface")
    logger.info("⚡ Backend: Python-based document conversion")
    logger.info("=" * 60)
    logger.info("🌐 Open your browser: http://localhost:5000")
    logger.info("=" * 60)

    try:
        # Use Werkzeug directly for more stable server
        from werkzeug.serving import run_simple
        run_simple('0.0.0.0', 5000, app, use_reloader=False, use_debugger=False, threaded=True)
    except KeyboardInterrupt:
        logger.info("Server stopped")
    except Exception as e:
        logger.error(f"Server error: {e}")
        import traceback
        traceback.print_exc()
