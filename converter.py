"""
Markdown to Word Document Converter
Converts markdown files to .docx format with full formatting support
Including: headings, callouts, images, tables, code blocks, and more
"""

import os
import re
import logging
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from PIL import Image
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Global session with connection pooling and retry strategy
_session_lock = threading.Lock()
_session = None

def get_session():
    """Get or create a requests session with optimized configuration"""
    global _session
    if _session is None:
        with _session_lock:
            if _session is None:
                _session = requests.Session()
                # Retry strategy for failed requests
                retry = Retry(
                    total=2,
                    backoff_factor=0.5,
                    status_forcelist=(500, 502, 504)
                )
                adapter = HTTPAdapter(max_retries=retry, pool_connections=10, pool_maxsize=10)
                _session.mount('http://', adapter)
                _session.mount('https://', adapter)
    return _session

class MarkdownConverter:
    """Converts markdown to Word documents"""

    # Image fetch configuration
    CORS_PROXIES = [
        'https://codetabs.com/code-api/cors?url=',
        'https://corsproxy.io/?',
        'https://proxy.cors.sh/',
        'https://cors.eu.org/?',
        'https://allorigins.win/raw?url='
    ]

    def __init__(self, markdown_file, output_file, include_images=True, max_workers=5,
                 add_page_numbers=True, add_page_border=True, add_watermark=True, add_logo=False):
        """Initialize converter with input/output paths"""
        self.markdown_file = markdown_file
        self.output_file = output_file
        self.include_images = include_images
        self.add_page_numbers = add_page_numbers
        self.add_page_border = add_page_border
        self.add_watermark = add_watermark
        self.add_logo = add_logo
        self.doc = Document()
        self.image_cache = {}
        self.max_workers = max_workers  # Thread pool size for concurrent image fetching
        self.pending_images = []  # Store image URLs to fetch concurrently
        self._setup_margins()

    def _setup_margins(self):
        """Set up document margins"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def convert(self):
        """Main conversion method"""
        try:
            with open(self.markdown_file, 'r', encoding='utf-8') as f:
                content = f.read()

            # PERFORMANCE: Extract and pre-fetch all images concurrently
            if self.include_images:
                image_urls = self._extract_all_image_urls(content)
                if image_urls:
                    self._prefetch_images(image_urls)

            lines = content.split('\n')
            i = 0

            while i < len(lines):
                line = lines[i]
                trimmed = line.strip()

                # Skip empty lines
                if not trimmed:
                    i += 1
                    continue

                # Skip HTML-only tags
                if trimmed in ['</br>', '<br>', '<br/>']:
                    i += 1
                    continue

                # Handle headings (##, ###, ####)
                if trimmed.startswith('## ') and not trimmed.startswith('### '):
                    text = trimmed[3:]
                    self.add_heading_1(text)
                    i += 1
                    continue

                if trimmed.startswith('#### '):
                    text = trimmed[5:]
                    self.add_heading_2(text)
                    i += 1
                    continue

                if trimmed.startswith('### '):
                    text = trimmed[4:]
                    self.add_heading_3(text)
                    i += 1
                    continue

                # Handle blockquotes (start with >)
                if trimmed.startswith('> '):
                    blockquote_text = ''
                    while i < len(lines) and lines[i].strip().startswith('> '):
                        blockquote_text += lines[i].strip()[2:] + ' '
                        i += 1
                    self.add_blockquote(blockquote_text.strip())
                    continue

                # Handle HTML tables
                if trimmed.startswith('<table'):
                    table_html = ''
                    while i < len(lines) and '</table>' not in lines[i]:
                        table_html += lines[i] + '\n'
                        i += 1
                    if i < len(lines):
                        table_html += lines[i]
                        i += 1
                    self._add_html_table(table_html)
                    continue

                # Handle markdown images ![alt](url)
                image_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', trimmed)
                if image_match:
                    alt_text = image_match.group(1)
                    image_url = image_match.group(2)
                    self.add_image_to_doc(image_url)
                    i += 1
                    continue

                # Handle images
                if '<img' in trimmed:
                    src_match = re.search(r'src="([^"]+)"', trimmed)
                    if src_match:
                        image_url = src_match.group(1)
                        self.add_image_to_doc(image_url)
                    i += 1
                    continue

                # Handle callouts
                if '**Note:**' in trimmed:
                    text = trimmed.replace('**Note:**', '').strip()
                    self.add_callout('note', text)
                    i += 1
                    continue

                if '**Tip:**' in trimmed:
                    text = trimmed.replace('**Tip:**', '').strip()
                    self.add_callout('tip', text)
                    i += 1
                    continue

                if '**Important:**' in trimmed:
                    text = trimmed.replace('**Important:**', '').strip()
                    self.add_callout('important', text)
                    i += 1
                    continue

                if '**Caution:**' in trimmed:
                    text = trimmed.replace('**Caution:**', '').strip()
                    self.add_callout('caution', text)
                    i += 1
                    continue

                if '**Congratulations:**' in trimmed:
                    text = trimmed.replace('**Congratulations:**', '').strip()
                    self.add_callout('congratulations', text)
                    i += 1
                    continue

                if '**Question:**' in trimmed:
                    text = trimmed.replace('**Question:**', '').strip()
                    self.add_callout('question', text)
                    i += 1
                    continue

                if '**Answer:**' in trimmed:
                    text = trimmed.replace('**Answer:**', '').strip()
                    self.add_callout('answer', text)
                    i += 1
                    continue

                # Skip empty numbered lines
                if re.match(r'^\d+\.\s*$', trimmed):
                    i += 1
                    continue

                # Handle code blocks
                if trimmed.startswith('```'):
                    code_content = ''
                    i += 1
                    while i < len(lines) and not lines[i].strip().startswith('```'):
                        code_content += lines[i] + '\n'
                        i += 1
                    i += 1  # Skip closing ```
                    self.add_code_block(code_content.rstrip())
                    continue

                # Handle bullet points
                if trimmed.startswith('- '):
                    text = trimmed[2:]
                    self.add_bullet_point(text)
                    i += 1
                    continue

                # Handle numbered steps (1. 1. format)
                step_match = re.match(r'^\d+\.\s+(\d+)\.\s+(.+)', trimmed)
                if step_match:
                    step_num = step_match.group(1)
                    step_text = step_match.group(2)
                    self.add_numbered_step(step_num, step_text)
                    i += 1
                    continue

                # Handle regular paragraphs
                if trimmed and not trimmed.startswith('<'):
                    self.add_paragraph(trimmed)
                    i += 1
                    continue

                i += 1

            # Apply document features before saving
            if self.add_page_numbers:
                self._add_page_numbers()

            if self.add_page_border:
                self._add_page_border()

            if self.add_watermark:
                self._add_watermark()

            if self.add_logo:
                self._add_logo_to_footer()

            # Save document
            self.doc.save(self.output_file)
            logger.info(f"✓ Document saved: {self.output_file}")
            return True

        except Exception as e:
            logger.error(f"Conversion error: {str(e)}")
            return False

    def add_heading_1(self, text):
        """Add Heading 1 (Dark Blue, 18pt)"""
        para = self.doc.add_paragraph(text)
        para.style = 'Heading 1'
        for run in para.runs:
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(30, 58, 138)  # Dark Blue #1e3a8a

        # Add bottom border
        pPr = para._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1e3a8a')
        pBdr.append(bottom)
        pPr.append(pBdr)

        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)

    def add_heading_2(self, text):
        """Add Heading 2 (Blue, 14pt, Bold)"""
        para = self.doc.add_paragraph(text)
        para.style = 'Heading 2'
        for run in para.runs:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(29, 78, 216)  # Blue #1d4ed8

        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)

    def add_heading_3(self, text):
        """Add Heading 3 (Light Blue, 13pt, Italic, Bold)"""
        para = self.doc.add_paragraph(text)
        para.style = 'Heading 3'
        for run in para.runs:
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.italic = True
            run.font.color.rgb = RGBColor(59, 130, 246)  # Light Blue #3b82f6

        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(4)

    def add_blockquote(self, text):
        """Add blockquote with attractive formatting and left border"""
        para = self.doc.add_paragraph()

        # Convert HTML tags to markdown and add formatted text
        text = self._convert_html_to_markdown(text)
        self._add_formatted_runs(para, text)

        # Set styling - reduced spacing for compact look
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.line_spacing = 1.15  # Reduced for compact appearance

        # Set background color (light grey)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'f0f0f0')  # Light grey background
        para._element.get_or_add_pPr().append(shading_elm)

        # Add left border
        pPr = para._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '24')  # Border width
        left.set(qn('w:space'), '0')
        left.set(qn('w:color'), '808080')  # Dark grey border
        pBdr.append(left)
        pPr.append(pBdr)

    def add_callout(self, callout_type, text):
        """Add callout box with specific styling and formatted text"""
        callout_config = {
            'note': {'color': RGBColor(37, 99, 235), 'bg': 'eff6ff', 'label': 'Note'},  # Blue
            'tip': {'color': RGBColor(22, 163, 74), 'bg': 'f0fdf4', 'label': 'Tip'},  # Green
            'important': {'color': RGBColor(220, 38, 38), 'bg': 'fef2f2', 'label': 'Important'},  # Red
            'caution': {'color': RGBColor(202, 138, 4), 'bg': 'fefce8', 'label': 'Caution'},  # Yellow
            'congratulations': {'color': RGBColor(147, 51, 234), 'bg': 'faf5ff', 'label': 'Congratulations'},  # Purple
            'question': {'color': RGBColor(146, 64, 14), 'bg': 'fef3c7', 'label': 'Question'},  # Brown
            'answer': {'color': RGBColor(22, 101, 52), 'bg': 'dcfce7', 'label': 'Answer'}  # Dark Green
        }

        config = callout_config.get(callout_type, callout_config['note'])

        para = self.doc.add_paragraph()

        # Add bold label
        label_run = para.add_run(f"{config['label']}: ")
        label_run.font.color.rgb = config['color']
        label_run.font.bold = True
        label_run.font.size = Pt(11)

        # Convert HTML tags to markdown and add formatted text
        text = self._convert_html_to_markdown(text)
        self._add_formatted_runs(para, text)

        # Set background color
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), config['bg'])
        para._element.get_or_add_pPr().append(shading_elm)

        # Add left border
        pPr = para._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '24')
        left.set(qn('w:space'), '0')
        # Convert RGB tuple to hex string
        rgb_color = config['color']
        hex_color = '%02x%02x%02x' % (rgb_color[0], rgb_color[1], rgb_color[2])
        left.set(qn('w:color'), hex_color)
        pBdr.append(left)
        pPr.append(pBdr)

        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)

    def add_bullet_point(self, text):
        """Add bullet point with formatting support"""
        para = self.doc.add_paragraph(style='List Bullet')
        self._add_formatted_runs(para, text)
        for run in para.runs:
            if run.font.size is None:
                run.font.size = Pt(12)
            if run.font.name != 'Consolas':
                run.font.name = 'Times New Roman'

        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)

    def add_numbered_step(self, step_num, text):
        """Add numbered step with formatting support and URL hyperlinks"""
        para = self.doc.add_paragraph()
        run_num = para.add_run(f"{step_num}. ")
        run_num.font.bold = True
        run_num.font.color.rgb = RGBColor(0, 0, 0)  # Black with bold
        run_num.font.size = Pt(12)
        run_num.font.name = 'Times New Roman'

        # Check if text contains URLs
        if 'http://' in text or 'https://' in text:
            # Parse and add with URL highlighting
            self._add_url_hyperlink(para, text)
        else:
            # Add step text with formatting support
            self._add_formatted_runs(para, text)

        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)

    def add_code_block(self, code):
        """Add code block with border, background, and proper formatting"""
        # Add label
        label_para = self.doc.add_paragraph()
        label_run = label_para.add_run('Code Block')
        label_run.font.bold = True
        label_run.font.size = Pt(10)
        label_run.font.color.rgb = RGBColor(75, 75, 75)
        label_para.paragraph_format.space_before = Pt(6)
        label_para.paragraph_format.space_after = Pt(2)

        # Add code with formatting
        para = self.doc.add_paragraph(style='Normal')
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.right_indent = Inches(0.5)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 0.9  # Very tight line spacing
        para.paragraph_format.line_spacing_rule = 1  # Exactly 0.9

        # Split code into lines and add with proper formatting
        code_lines = code.split('\n')
        for idx, line in enumerate(code_lines):
            if idx > 0:
                run_newline = para.add_run('\n')
                run_newline.font.size = Pt(1)  # Minimal size for newline
            run = para.add_run(line if line else ' ')
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Set background color for entire code block
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'f3f4f6')  # Light gray
        para._element.get_or_add_pPr().append(shading_elm)

        # Add border around code block
        pPr = para._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'cccccc')
            pBdr.append(border)
        pPr.append(pBdr)

        para.paragraph_format.space_after = Pt(6)

    def _add_url_hyperlink(self, paragraph, text):
        """Add text with URL hyperlinks as blue bold clickable links"""
        # Convert HTML tags to markdown first
        text = self._convert_html_to_markdown(text)

        # Pattern to match URLs - stop at space or common punctuation, including backticks
        url_pattern = r'(https?://[^\s),\.\!\?;:`]+)'
        parts = re.split(url_pattern, text)

        for part in parts:
            if re.match(r'https?://', part):
                # This is a URL - add as clickable hyperlink
                # Remove any backticks around URL
                clean_part = part.strip('`')
                run = paragraph.add_run(clean_part)
                run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
                run.font.bold = True
                run.font.size = Pt(12)
                run.font.underline = True
                run.font.name = 'Times New Roman'

                # Add hyperlink relationship
                try:
                    part_elem = run._element
                    # Get relationship ID from parent document
                    rel_id = paragraph._parent._element.getparent().part.relate_to(
                        clean_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                        is_external=True
                    )
                except Exception as e:
                    logger.debug(f"Could not create hyperlink for {clean_part}: {str(e)}")
            else:
                # Regular text with formatting support
                if part:
                    # Remove backticks from regular text too
                    part = part.replace('`', '')
                    if part:
                        self._add_formatted_runs(paragraph, part)

    def _convert_html_to_markdown(self, text):
        """Convert HTML formatting tags to markdown for consistent parsing"""
        # First convert HTML entities
        text = text.replace('&nbsp;', ' ')
        text = text.replace('&lt;', '<')
        text = text.replace('&gt;', '>')
        text = text.replace('&amp;', '&')

        # Convert <strong> to ** - match content between tags, handle attributes
        text = re.sub(r'<strong[^>]*>([^<]*)</strong>', r'**\1**', text, flags=re.IGNORECASE)
        # Convert <b> to **
        text = re.sub(r'<b[^>]*>([^<]*)</b>', r'**\1**', text, flags=re.IGNORECASE)
        # Convert <em> to *
        text = re.sub(r'<em[^>]*>([^<]*)</em>', r'*\1*', text, flags=re.IGNORECASE)
        # Convert <i> to *
        text = re.sub(r'<i[^>]*>([^<]*)</i>', r'*\1*', text, flags=re.IGNORECASE)
        # Convert <code> to backticks
        text = re.sub(r'<code[^>]*>([^<]*)</code>', r'`\1`', text, flags=re.IGNORECASE)
        # Remove any remaining HTML tags that weren't converted
        text = re.sub(r'<[^>]+>', '', text)
        return text

    def add_paragraph(self, text):
        """Add regular paragraph with formatting support"""
        para = self.doc.add_paragraph()
        text = self._convert_html_to_markdown(text)
        self._add_formatted_runs(para, text)
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(4)

    def _add_formatted_runs(self, paragraph, text):
        """Parse and add formatted text runs - exactly matching example.html parseInlineFormatting"""
        remaining = text

        # Replace patterns with markers - EXACT ORDER matching example.html
        # Use proper backtick pattern with start/end markers
        remaining = re.sub(r'`([^`]+)`', r'<<CODE>>\1<<ENDCODE>>', remaining)
        # Handle bold (**text**) - non-greedy
        remaining = re.sub(r'\*\*(.+?)\*\*', r'<<BOLD>>\1<<ENDBOLD>>', remaining)
        # Handle italic (*text*) - non-greedy
        remaining = re.sub(r'\*(.+?)\*', r'<<ITALIC>>\1<<ENDITALIC>>', remaining)

        # Also handle HTML tags
        remaining = re.sub(r'<strong>([^<]+)</strong>', r'<<BOLD>>\1<<ENDBOLD>>', remaining)
        remaining = re.sub(r'<em>([^<]+)</em>', r'<<ITALIC>>\1<<ENDITALIC>>', remaining)

        # Split by all markers using capture group to keep markers
        parts = re.split(r'(<<BOLD>>|<<ENDBOLD>>|<<ITALIC>>|<<ENDITALIC>>|<<CODE>>|<<ENDCODE>>)', remaining)

        is_bold = False
        is_italic = False
        is_code = False

        for part in parts:
            if not part:
                continue

            if part == '<<BOLD>>':
                is_bold = True
            elif part == '<<ENDBOLD>>':
                is_bold = False
            elif part == '<<ITALIC>>':
                is_italic = True
            elif part == '<<ENDITALIC>>':
                is_italic = False
            elif part == '<<CODE>>':
                is_code = True
            elif part == '<<ENDCODE>>':
                is_code = False
            else:
                # Add text run with appropriate formatting
                if part:  # Skip empty parts
                    run = paragraph.add_run(part)
                    run.font.bold = is_bold
                    run.font.italic = is_italic
                    if is_code:
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    else:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)

    def fetch_image_with_cors(self, url):
        """Fetch image with CORS proxy fallback using optimized session"""
        session = get_session()

        # Try direct fetch first
        try:
            response = session.get(url, timeout=10)
            response.raise_for_status()
            logger.info(f"✓ Image fetched (direct): {url[-50:]} ({len(response.content)} bytes)")
            return response.content
        except Exception as e:
            logger.debug(f"Direct fetch failed: {str(e)}")

        # Try CORS proxies
        for proxy in self.CORS_PROXIES:
            try:
                proxy_url = proxy + url
                response = session.get(proxy_url, timeout=10)
                response.raise_for_status()
                logger.info(f"✓ Image fetched (proxy): {url[-30:]} ({len(response.content)} bytes)")
                return response.content
            except Exception as e:
                logger.debug(f"Proxy failed: {proxy}: {str(e)}")
                continue

        logger.error(f"✗ Could not fetch image: {url}")
        return None

    def add_image_to_doc(self, image_url):
        """Fetch and add image to document"""
        if not self.include_images:
            logger.info(f"Skipping image (include_images=False): {image_url}")
            return

        # Add spacing before image for better presentation
        spacing_para = self.doc.add_paragraph()
        spacing_para.paragraph_format.space_before = Pt(6)
        spacing_para.paragraph_format.space_after = Pt(6)
        spacing_para.paragraph_format.line_spacing = 1.5

        # Check cache
        if image_url in self.image_cache:
            image_data = self.image_cache[image_url]
        else:
            image_data = self.fetch_image_with_cors(image_url)
            if not image_data:
                logger.warning(f"Failed to fetch image: {image_url}")
                return
            self.image_cache[image_url] = image_data

        try:
            # Optimize image
            img = Image.open(BytesIO(image_data))
            if img.mode == 'RGBA':
                img = img.convert('RGB')

            # Save optimized to BytesIO
            optimized = BytesIO()
            img.save(optimized, format='JPEG', quality=85)
            optimized.seek(0)

            # Add to document - CRITICAL: Use doc.add_picture, NOT run.add_picture()
            max_width = 5.5
            picture_para = self.doc.add_picture(optimized, width=Inches(max_width))

            logger.info(f"✓ Image embedded successfully: {image_url[-30:]}")

        except Exception as e:
            logger.error(f"Failed to embed image: {str(e)}")

    def _add_html_table(self, html):
        """Parse and add HTML table to document with proper formatting"""
        try:
            # Extract rows
            rows = re.findall(r'<tr[^>]*>(.*?)</tr>', html, re.DOTALL)
            if not rows:
                return

            # Extract cells from first row to determine column count
            first_row_cells = re.findall(r'<t[dh][^>]*>(.*?)</t[dh]>', rows[0], re.DOTALL)
            col_count = len(first_row_cells)

            # Create table
            table = self.doc.add_table(rows=len(rows), cols=col_count)
            table.style = 'Light Grid Accent 1'

            # Fill table with proper formatting
            for row_idx, row_html in enumerate(rows):
                # Check if this is a header row (contains <th> tags)
                is_header_row = '<th' in row_html.lower()

                # Extract cells - handle both <th> and <td>
                cells = re.findall(r'<t[dh][^>]*>(.*?)</t[dh]>', row_html, re.DOTALL)
                for cell_idx, cell_content in enumerate(cells[:col_count]):
                    cell = table.rows[row_idx].cells[cell_idx]
                    cell_content = cell_content.strip()

                    # Add content to cell with formatting support
                    if cell_content:
                        # Convert HTML tags to markdown for proper formatting
                        cell_content = self._convert_html_to_markdown(cell_content)
                        # Clean up extra whitespace
                        cell_content = cell_content.strip()

                        # Clear default paragraph
                        cell.text = ''
                        para = cell.paragraphs[0]

                        # If header row, make the text bold
                        if is_header_row:
                            self._add_formatted_runs(para, cell_content)
                            # Apply bold to all runs in header
                            for run in para.runs:
                                run.font.bold = True
                                run.font.size = Pt(12)
                                if run.font.name != 'Consolas':
                                    run.font.name = 'Times New Roman'
                        else:
                            self._add_formatted_runs(para, cell_content)
                            for run in para.runs:
                                if run.font.size is None:
                                    run.font.size = Pt(12)
                                if run.font.name != 'Consolas':
                                    run.font.name = 'Times New Roman'

                        para.paragraph_format.space_before = Pt(2)
                        para.paragraph_format.space_after = Pt(2)

        except Exception as e:
            logger.error(f"Failed to parse table: {str(e)}")

    def _prefetch_images(self, image_urls):
        """Pre-fetch multiple images concurrently"""
        if not self.include_images or not image_urls:
            return

        unique_urls = list(set(image_urls))  # Remove duplicates
        logger.info(f"Pre-fetching {len(unique_urls)} unique images concurrently...")

        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            future_to_url = {
                executor.submit(self.fetch_image_with_cors, url): url
                for url in unique_urls
                if url not in self.image_cache
            }

            completed = 0
            for future in as_completed(future_to_url):
                url = future_to_url[future]
                try:
                    image_data = future.result(timeout=15)
                    if image_data:
                        self.image_cache[url] = image_data
                        completed += 1
                except Exception as e:
                    logger.warning(f"Failed to pre-fetch image: {url[-50:]}: {str(e)}")

            logger.info(f"✓ Pre-fetched {completed}/{len(unique_urls)} images")

    def _extract_all_image_urls(self, content):
        """Extract all image URLs from markdown content"""
        urls = []
        # Extract HTML img tags
        for match in re.finditer(r'<img[^>]+src="([^"]+)"', content):
            urls.append(match.group(1))
        # Extract markdown images ![alt](url)
        for match in re.finditer(r'!\[([^\]]*)\]\(([^)]+)\)', content):
            urls.append(match.group(2))
        return urls

    def to_html(self):
        """Generate HTML preview that matches conversion output"""
        try:
            with open(self.markdown_file, 'r', encoding='utf-8') as f:
                content = f.read()

            # Extract and cache images for better performance
            if self.include_images:
                image_urls = self._extract_all_image_urls(content)
                if image_urls:
                    self._prefetch_images(image_urls)

            lines = content.split('\n')
            html = '<div style="font-family: \'Times New Roman\', serif; max-width: 8.5in; color: #333; line-height: 1.6; background: white; padding: 20px;">'
            i = 0

            while i < len(lines):
                line = lines[i]
                trimmed = line.strip()

                # Skip empty lines
                if not trimmed:
                    i += 1
                    continue

                # Skip HTML-only tags
                if trimmed in ['</br>', '<br>', '<br/>']:
                    i += 1
                    continue

                # Handle headings
                if trimmed.startswith('## ') and not trimmed.startswith('### '):
                    text = self._escape_html(trimmed[3:])
                    html += f'<h2 style="color: #1e3a8a; font-size: 18pt; font-weight: bold; margin: 12pt 0 6pt 0; border-bottom: 2px solid #1e3a8a; padding-bottom: 4px;">{text}</h2>'
                    i += 1
                    continue

                if trimmed.startswith('#### '):
                    text = self._escape_html(trimmed[5:])
                    html += f'<h3 style="color: #1d4ed8; font-size: 14pt; font-weight: bold; margin: 10pt 0 4pt 0;">{text}</h3>'
                    i += 1
                    continue

                if trimmed.startswith('### '):
                    text = self._escape_html(trimmed[4:])
                    html += f'<h4 style="color: #3b82f6; font-size: 13pt; font-weight: bold; font-style: italic; margin: 8pt 0 4pt 0;">{text}</h4>'
                    i += 1
                    continue

                # Handle HTML tables
                if trimmed.startswith('<table'):
                    table_html = ''
                    while i < len(lines) and '</table>' not in lines[i]:
                        table_html += lines[i] + '\n'
                        i += 1
                    if i < len(lines):
                        table_html += lines[i]
                        i += 1
                    html += self._render_html_table(table_html)
                    continue

                # Handle markdown images ![alt](url)
                image_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', trimmed)
                if image_match:
                    alt_text = image_match.group(1)
                    image_url = image_match.group(2)
                    html += f'<div style="margin: 12px 0;"><img src="{self._escape_html(image_url)}" alt="{self._escape_html(alt_text)}" style="max-width: 100%; height: auto; border: 1px solid #ddd; border-radius: 4px;" /></div>'
                    i += 1
                    continue

                # Handle images
                if '<img' in trimmed:
                    src_match = re.search(r'src="([^"]+)"', trimmed)
                    if src_match:
                        image_url = src_match.group(1)
                        html += f'<div style="margin: 12px 0;"><img src="{self._escape_html(image_url)}" alt="Image" style="max-width: 100%; height: auto; border: 1px solid #ddd; border-radius: 4px;" /></div>'
                    i += 1
                    continue

                # Handle callouts
                callout_types = {
                    '**Note:**': ('note', '#eff6ff', '#2563eb', '#1e40af'),
                    '**Tip:**': ('tip', '#f0fdf4', '#16a34a', '#15803d'),
                    '**Important:**': ('important', '#fef2f2', '#dc2626', '#991b1b'),
                    '**Caution:**': ('caution', '#fefce8', '#ca8a04', '#92400e'),
                    '**Congratulations:**': ('congratulations', '#faf5ff', '#9333ea', '#6b21a8'),
                }

                found_callout = False
                for prefix, (_, bg_color, border_color, text_color) in callout_types.items():
                    if trimmed.startswith(prefix):
                        raw_text = trimmed[len(prefix):].strip()
                        # Apply inline formatting (bold, italic, code) before adding to HTML
                        text = self._format_inline_html(raw_text)
                        label = prefix.replace('**', '').replace(':', '')
                        html += f'<div style="background: {bg_color}; border-left: 4px solid {border_color}; padding: 8px 12px; margin: 6pt 0; border-radius: 4px;"><span style="color: {border_color}; font-weight: bold;">{label}:</span> <span style="color: {text_color};">{text}</span></div>'
                        found_callout = True
                        i += 1
                        break

                if found_callout:
                    continue

                # Handle code blocks
                if trimmed.startswith('```'):
                    code_content = ''
                    i += 1
                    while i < len(lines) and not lines[i].strip().startswith('```'):
                        code_content += self._escape_html(lines[i]) + '\n'
                        i += 1
                    html += f'<div style="background: #f3f4f6; border: 1px solid #cccccc; border-radius: 4px; padding: 12px; margin: 6pt 0;"><div style="font-size: 9pt; color: #4b5563; margin-bottom: 4px; font-weight: bold;">Code Block</div><pre style="font-family: Consolas, monospace; font-size: 9pt; color: #000; margin: 0; white-space: pre-wrap; word-wrap: break-word;">{code_content}</pre></div>'
                    i += 1
                    continue

                # Handle bullet points
                if trimmed.startswith('- '):
                    text = self._format_inline_html(trimmed[2:])
                    html += f'<div style="margin-left: 20px; margin: 2pt 0;">• {text}</div>'
                    i += 1
                    continue

                # Handle numbered steps (1. 1. format)
                step_match = re.match(r'^\d+\.\s+(\d+)\.\s+(.+)', trimmed)
                if step_match:
                    step_num = step_match.group(1)
                    step_text = self._format_inline_html(step_match.group(2))
                    html += f'<div style="margin-left: 20px; margin: 4pt 0;"><span style="font-weight: bold;">{step_num}.</span> {step_text}</div>'
                    i += 1
                    continue

                # Regular paragraph
                text = self._format_inline_html(trimmed)
                html += f'<p style="margin: 4pt 0; line-height: 1.3;">{text}</p>'
                i += 1

            html += '</div>'
            return html

        except Exception as e:
            logger.error(f"Failed to generate HTML preview: {str(e)}")
            return f'<div style="color: red; padding: 20px;">Error generating preview: {str(e)}</div>'

    def _escape_html(self, text):
        """Escape HTML special characters"""
        if not text:
            return ''
        text = str(text)
        return (text.replace('&', '&amp;')
                    .replace('<', '&lt;')
                    .replace('>', '&gt;')
                    .replace('"', '&quot;')
                    .replace("'", '&#39;'))

    def _format_inline_html(self, text):
        """Format inline markdown to HTML"""
        text = self._escape_html(text)
        # Bold
        text = re.sub(r'\*\*(.+?)\*\*', r'<strong style="font-weight: bold;">\1</strong>', text)
        # Italic
        text = re.sub(r'\*(.+?)\*', r'<em style="font-style: italic;">\1</em>', text)
        # Inline code
        text = re.sub(r'`([^`]+)`', r'<code style="background: #f0f0f0; padding: 2px 4px; font-family: Consolas, monospace; font-size: 9pt; color: #000;">\1</code>', text)
        return text

    def _render_html_table(self, table_html):
        """Render HTML table with Word-style formatting"""
        try:
            # Parse table
            table_html_stripped = table_html.strip()

            if not table_html_stripped:
                return ''

            html_output = '<table style="width: 100%; border-collapse: collapse; margin: 12px 0; border: 1px solid #bbb;">'

            # Extract rows with more robust regex
            rows = re.findall(r'<tr[^>]*>(.*?)</tr>', table_html_stripped, re.DOTALL)

            if not rows:
                return '<div style="color: #666; padding: 8px; background: #f9f9f9; border-radius: 4px;">No table data found</div>'

            for row_idx, row in enumerate(rows):
                bg_color = '#e6e6e6' if '<th' in row else ('#ffffff' if row_idx % 2 == 0 else '#f9f9f9')
                html_output += f'<tr style="background: {bg_color};">'

                # Extract cells - use non-greedy matching to avoid overlapping groups
                cell_pattern = r'<(?:th|td)[^>]*>(.*?)</(?:th|td)>'
                cells = re.findall(cell_pattern, row, re.DOTALL)

                if not cells:
                    # Fallback: split by common delimiters
                    cell_contents = re.split(r'(?:<[^>]+>)+', row.strip())
                    cells = [content.strip() for content in cell_contents if content.strip()]

                for cell_content in cells:
                    cell_style = 'border: 1px solid #bbb; padding: 8px; color: #333;'
                    # Remove all HTML tags for cleaner text
                    cell_text = self._escape_html(re.sub(r'<[^>]+>', '', str(cell_content)).strip())
                    html_output += f'<td style="{cell_style}">{cell_text}</td>'

                html_output += '</tr>'

            html_output += '</table>'
            return html_output
        except Exception as e:
            logger.error(f"Failed to render table: {str(e)}")
            return f'<div style="color: #666; padding: 8px; background: #f9f9f9; border-radius: 4px;">Table rendering fixed with fallback</div>'

    def _add_page_numbers(self):
        """Add page numbers to footer (centered at bottom)"""
        try:
            for section in self.doc.sections:
                footer = section.footer

                # Check if footer already has content (logo)
                if footer.paragraphs:
                    # Use first paragraph for page numbers
                    footer_para = footer.paragraphs[0]
                    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Changed to CENTER
                else:
                    footer_para = footer.add_paragraph()
                    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Changed to CENTER

                # No text, just the field code for page number
                footer_para.text = ''

                # Create page number field
                run = footer_para.add_run()

                # Use WordprocessingML field code for page numbers
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')

                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = 'PAGE'

                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'end')

                run._r.append(fldChar1)
                run._r.append(instrText)
                run._r.append(fldChar2)

                # Style the page number - position at bottom
                footer_para.runs[0].font.size = Pt(10)
                footer_para.paragraph_format.space_before = Pt(6)
                footer_para.paragraph_format.space_after = Pt(0)  # No space after to keep at bottom

            logger.info("✓ Page numbers added to footer (centered at bottom)")
        except Exception as e:
            logger.error(f"Failed to add page numbers: {str(e)}")

    def _add_page_border(self):
        """Add thick page borders with dark blue color #00264d"""
        try:
            for section in self.doc.sections:
                sectPr = section._sectPr

                # Remove existing borders
                existing_borders = sectPr.findall(qn('w:pgBorders'))
                for border in existing_borders:
                    sectPr.remove(border)

                # Create page borders element
                pgBorders = OxmlElement('w:pgBorders')
                pgBorders.set(qn('w:offsetFrom'), 'page')

                # Dark blue color (#00264d)
                border_color = '00264d'  # Dark blue
                border_size = '48'  # Extra thick (triple line equivalent)

                # Create thick borders for all sides
                for border_name in ['top', 'bottom', 'left', 'right']:
                    border_el = OxmlElement(f'w:{border_name}')
                    border_el.set(qn('w:val'), 'triple')  # Triple line style for thickness
                    border_el.set(qn('w:sz'), border_size)  # Size in 1/8 pt
                    border_el.set(qn('w:space'), '24')  # Space from content
                    border_el.set(qn('w:color'), border_color)
                    border_el.set(qn('w:shadow'), '0')
                    pgBorders.append(border_el)

                sectPr.append(pgBorders)

            logger.info("✓ Page borders added (triple line, dark blue #00264d, thick)")
        except Exception as e:
            logger.error(f"Failed to add page borders: {str(e)}")

    def _add_watermark(self):
        """Add diagonal text watermark 'XtremeLabs LLC' with specified properties to headers"""
        try:
            from lxml import etree

            # VML Shape for the watermark
            vml_script = '''<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:rPr><w:noProof/></w:rPr>
                <w:pict>
                    <v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136" adj="10800" path="m@7,l@8,m@5,21600l@6,21600e" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office">
                        <v:formulas>
                            <v:f eqn="sum #0 0 10800"/><v:f eqn="prod #0 2 1"/><v:f eqn="sum 21600 0 @1"/><v:f eqn="sum 0 0 @2"/><v:f eqn="sum 21600 0 @3"/><v:f eqn="if @0 @3 0"/><v:f eqn="if @0 21600 @1"/><v:f eqn="if @0 0 @2"/><v:f eqn="if @0 @4 21600"/><v:f eqn="mid @5 @6"/><v:f eqn="mid @8 @5"/><v:f eqn="mid @7 @8"/><v:f eqn="mid @6 @7"/><v:f eqn="sum @6 0 @5"/>
                        </v:formulas>
                        <v:path textpathok="t" o:connecttype="custom" o:connectlocs="@9,0;@10,10800;@11,21600;@12,10800" o:connectangles="270,180,90,0"/>
                        <v:textpath on="t" fitshape="t"/>
                        <o:lock v:ext="edit" text="t" shapetype="t"/>
                    </v:shapetype>
                    <v:shape id="PowerPlusWaterMarkObject" o:spid="_x0000_s102" type="#_x0000_t136" style="position:absolute;left:0;text-align:left;margin-left:0;margin-top:0;width:500pt;height:250pt;rotation:315;z-index:-251658240;mso-position-horizontal:center;mso-position-horizontal-relative:margin;mso-position-vertical:center;mso-position-vertical-relative:margin" o:allowincell="f" fillcolor="#00264d" stroked="f" xmlns:v="urn:schemas-microsoft-com:vml" xmlns:w10="urn:schemas-microsoft-com:office:word" xmlns:o="urn:schemas-microsoft-com:office:office">
                        <v:fill opacity=".15"/>
                        <v:textpath style="font-family:'Cambria';font-size:1pt;font-weight:bold" string="XtremeLabs LLC"/>
                        <w10:wrap anchorx="margin" anchory="margin"/>
                    </v:shape>
                </w:pict>
            </w:r>'''

            pict = etree.fromstring(vml_script)

            for section in self.doc.sections:
                header = section.header
                # Get the first paragraph or create one
                if not header.paragraphs:
                    para = header.add_paragraph()
                else:
                    para = header.paragraphs[0]
                
                # Append the watermark picture element directly to the paragraph
                import copy
                para._p.append(copy.deepcopy(pict))

            logger.info("✓ Text watermark 'XtremeLabs LLC' added to headers (diagonal, Cambria font, dark blue #00264d)")
        except Exception as e:
            logger.error(f"Failed to add watermark: {str(e)}")

    def _add_logo_to_footer(self):
        """Add logo image to footer (bottom right corner) - works with centered page numbers"""
        try:
            logo_path = os.path.join(os.path.dirname(__file__), 'xtremelabs logo.jpg')

            if not os.path.exists(logo_path):
                logger.warning(f"Logo file not found at {logo_path}")
                return

            for section in self.doc.sections:
                footer = section.footer

                # Check if footer already has a logo
                has_logo = any('xtremelabs logo.jpg' in str(para._element.xml) or
                              'pic:' in str(para._element.xml) for para in footer.paragraphs)
                if has_logo:
                    logger.info("✓ Logo already exists in footer, skipping addition")
                    continue

                # Check if footer already has page numbers
                has_page_numbers = any('PAGE' in str(para._element.xml) for para in footer.paragraphs)

                if has_page_numbers and footer.paragraphs:
                    # Page numbers are centered, add logo to a new paragraph
                    logo_para = footer.add_paragraph()
                else:
                    # No page numbers, add logo alone
                    if footer.paragraphs:
                        logo_para = footer.paragraphs[0]
                    else:
                        logo_para = footer.add_paragraph()

                # Clear only the current paragraph if it's empty
                logo_para.text = ''
                logo_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                logo_para.paragraph_format.space_before = Pt(6)  # Consistent spacing
                logo_para.paragraph_format.space_after = Pt(0)   # Keep at bottom
                logo_para.paragraph_format.right_indent = Inches(0.1)  # Small indent for bottom right positioning

                # Add logo image
                try:
                    run = logo_para.add_run()
                    run.add_picture(logo_path, width=Inches(0.5))  # Smaller size
                    logger.info(f"✓ Logo added to footer from: {logo_path}")
                except Exception as e:
                    logger.warning(f"Could not add logo image: {str(e)}")

            logger.info("✓ Logo added to footer (bottom right, smaller size)")
        except Exception as e:
            logger.error(f"Failed to add logo to footer: {str(e)}")


if __name__ == '__main__':
    import sys

    if len(sys.argv) < 3:
        print("Usage: python converter.py <input.md> <output.docx>")
        sys.exit(1)

    converter = MarkdownConverter(sys.argv[1], sys.argv[2])
    success = converter.convert()
    sys.exit(0 if success else 1)
